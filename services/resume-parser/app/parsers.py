"""PDF and DOCX text extraction."""
from __future__ import annotations

import io
from typing import Tuple


def extract_pdf(content: bytes) -> Tuple[str, list[str]]:
    """
    Extract text from a native (machine-generated) PDF.
    Returns (full_text, warnings).
    """
    try:
        import pdfplumber
    except ImportError:
        return "", ["pdfplumber_not_installed"]

    warnings: list[str] = []
    pages_text: list[str] = []

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if text:
                    pages_text.append(text)
                else:
                    warnings.append(f"page_{page.page_number}_no_text_extracted")
    except Exception as exc:
        warnings.append(f"pdf_parse_error:{exc}")

    full_text = "\n".join(pages_text)

    # Heuristic: if very little text extracted and file is large → likely scanned
    if len(full_text.strip()) < 200 and len(content) > 50_000:
        warnings.append("likely_scanned_resume_ocr_fallback_needed")

    return full_text, warnings


def extract_docx(content: bytes) -> Tuple[str, list[str]]:
    """Extract text from a DOCX file. Returns (full_text, warnings)."""
    try:
        from docx import Document
    except ImportError:
        return "", ["python_docx_not_installed"]

    warnings: list[str] = []
    paragraphs: list[str] = []

    try:
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also pull text from tables (common for skills sections)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
                    if not any("table" in w for w in warnings):
                        warnings.append("table_content_detected_may_affect_ats")
    except Exception as exc:
        warnings.append(f"docx_parse_error:{exc}")

    return "\n".join(paragraphs), warnings


def ocr_pdf(content: bytes) -> Tuple[str, list[str]]:
    """OCR a scanned/image-only PDF. Returns (text, warnings).

    Renders each page to an image and runs Tesseract. Degrades honestly: if the
    OCR libraries or the Tesseract binary are unavailable, returns an empty
    string with a specific warning rather than raising.
    """
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        return "", ["ocr_unavailable_dependencies_missing"]

    warnings: list[str] = []
    try:
        images = convert_from_bytes(content, dpi=200)
    except Exception as exc:  # pdf2image needs poppler; surface why, don't crash
        return "", [f"ocr_render_failed:{type(exc).__name__}"]

    pages_text: list[str] = []
    for page_no, image in enumerate(images, start=1):
        try:
            page_text = pytesseract.image_to_string(image)
        except Exception as exc:
            warnings.append(f"ocr_page_{page_no}_failed:{type(exc).__name__}")
            continue
        if page_text.strip():
            pages_text.append(page_text)

    text = "\n".join(pages_text)
    if text.strip():
        warnings.append("ocr_applied_scanned_pdf")
    return text, warnings
